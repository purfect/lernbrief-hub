from __future__ import annotations

import os
import sqlite3
import random
import re
import sys
import socket
import json
import shutil
import threading
import time
import webbrowser
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from flask import Flask, flash, g, redirect, render_template, request, send_file, url_for

BASE_DIR = Path(__file__).resolve().parent
if getattr(sys, "frozen", False):
    RESOURCE_DIR = Path(getattr(sys, "_MEIPASS", BASE_DIR))
    DATA_DIR = Path(sys.executable).resolve().parent
else:
    RESOURCE_DIR = BASE_DIR
    DATA_DIR = BASE_DIR

DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / "lernbrief_hub.db"

GRADE_OPTIONS = [1, 2, 3, 4, 5, 6]
SEMESTER_PATTERN = re.compile(r"^(\d{4})/(\d{4})-(HJ[12])$")


def current_school_semester(now: datetime | None = None) -> str:
    """Return current German school semester key as YYYY/YYYY-HJ1|HJ2.

    HJ1 runs from August to January and HJ2 from February to July.
    """
    dt = now or datetime.now()
    year = dt.year
    month = dt.month

    if month >= 8:
        start_year = year
        end_year = year + 1
        half = "HJ1"
    elif month == 1:
        start_year = year - 1
        end_year = year
        half = "HJ1"
    else:
        start_year = year - 1
        end_year = year
        half = "HJ2"

    return f"{start_year}/{end_year}-{half}"


def parse_semester(value: str) -> tuple[int, int, str] | None:
    match = SEMESTER_PATTERN.match(value)
    if not match:
        return None
    start_year = int(match.group(1))
    end_year = int(match.group(2))
    half = match.group(3)
    if end_year != start_year + 1:
        return None
    return start_year, end_year, half


def semester_sort_key(value: str) -> tuple[int, int]:
    parsed = parse_semester(value)
    if parsed is None:
        return (-1, -1)
    start_year, _, half = parsed
    half_idx = 1 if half == "HJ1" else 2
    return (start_year, half_idx)


def school_semester_options(
    db: sqlite3.Connection | None = None,
    anchor: datetime | None = None,
    years_back: int | None = None,
    years_forward: int | None = None,
) -> list[str]:
    """Build selectable German school semesters around current year plus used data years."""
    dt = anchor or datetime.now()
    current = current_school_semester(dt)
    school_year_start = int(current.split("/")[0])

    back = years_back if years_back is not None else int(os.getenv("SEMESTER_YEARS_BACK", "3"))
    forward = years_forward if years_forward is not None else int(os.getenv("SEMESTER_YEARS_FORWARD", "3"))

    options: set[str] = set()
    for start_year in range(school_year_start - back, school_year_start + forward + 1):
        end_year = start_year + 1
        options.add(f"{start_year}/{end_year}-HJ1")
        options.add(f"{start_year}/{end_year}-HJ2")

    if db is not None:
        used_rows = db.execute(
            """
            SELECT semester FROM ratings
            UNION
            SELECT semester FROM letters
            """
        ).fetchall()
        for row in used_rows:
            sem = (row["semester"] or "").strip()
            if parse_semester(sem):
                options.add(sem)

    return sorted(options, key=semester_sort_key, reverse=True)


def normalize_semester(raw_value: str | None) -> str:
    value = (raw_value or "").strip()
    if parse_semester(value):
        return value
    return DEFAULT_SEMESTER


def ensure_sentence_punctuation(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return cleaned
    if cleaned[-1] not in ".!?":
        return f"{cleaned}."
    return cleaned


def html_to_plain_text(raw_html: str) -> str:
    text = raw_html or ""
    text = re.sub(r"(?i)<\s*br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</\s*p\s*>", "\n", text)
    text = re.sub(r"(?i)<\s*/\s*div\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def safe_format_text(template_text: str, values: dict[str, Any]) -> str:
    class SafeTemplateValues(dict[str, Any]):
        def __missing__(self, key: str) -> str:
            return "{" + key + "}"

    return template_text.format_map(SafeTemplateValues(values))


def map_export_font(font_family: str) -> tuple[str, str]:
    family = (font_family or "").lower()
    if "times" in family or "georgia" in family:
        return "Times-Roman", "Times New Roman"
    if "courier" in family or "mono" in family:
        return "Courier", "Courier New"
    return "Helvetica", "Calibri"


def parse_css_style(style_text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for chunk in (style_text or "").split(";"):
        if ":" not in chunk:
            continue
        k, v = chunk.split(":", 1)
        out[k.strip().lower()] = v.strip()
    return out


def parse_font_size_to_pt(raw: str) -> float | None:
    val = (raw or "").strip().lower()
    if not val:
        return None
    try:
        if val.endswith("pt"):
            return float(val[:-2])
        if val.endswith("px"):
            return round(float(val[:-2]) * 0.75, 1)
        return float(val)
    except ValueError:
        return None


class RichTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.paragraphs: list[list[dict[str, Any]]] = [[]]
        self.style_stack: list[dict[str, Any]] = [{"bold": False, "italic": False, "underline": False, "font": None, "size": None}]

    @property
    def style(self) -> dict[str, Any]:
        return self.style_stack[-1]

    def _push_style(self, **updates: Any) -> None:
        next_style = dict(self.style)
        next_style.update(updates)
        self.style_stack.append(next_style)

    def _pop_style(self) -> None:
        if len(self.style_stack) > 1:
            self.style_stack.pop()

    def _ensure_paragraph(self) -> None:
        if not self.paragraphs:
            self.paragraphs.append([])

    def _new_paragraph(self) -> None:
        if self.paragraphs and not self.paragraphs[-1]:
            return
        self.paragraphs.append([])

    def _append_text(self, text: str) -> None:
        if not text:
            return
        self._ensure_paragraph()
        self.paragraphs[-1].append({"text": text, **dict(self.style)})

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attrs_dict = {k.lower(): (v or "") for k, v in attrs}

        if tag in {"p", "div"}:
            self._ensure_paragraph()
            return
        if tag == "br":
            self._append_text("\n")
            return
        if tag in {"b", "strong"}:
            self._push_style(bold=True)
            return
        if tag in {"i", "em"}:
            self._push_style(italic=True)
            return
        if tag == "u":
            self._push_style(underline=True)
            return
        if tag in {"span", "font"}:
            updates: dict[str, Any] = {}
            style_map = parse_css_style(attrs_dict.get("style", ""))
            if "font-family" in style_map:
                updates["font"] = style_map["font-family"].split(",")[0].strip("'\"")
            if "font-size" in style_map:
                updates["size"] = parse_font_size_to_pt(style_map["font-size"])
            if "face" in attrs_dict:
                updates["font"] = attrs_dict["face"].split(",")[0].strip("'\"")
            if "size" in attrs_dict:
                updates["size"] = parse_font_size_to_pt(attrs_dict["size"])
            self._push_style(**updates)
            return

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"b", "strong", "i", "em", "u", "span", "font"}:
            self._pop_style()
        if tag in {"p", "div"}:
            self._new_paragraph()

    def handle_data(self, data: str) -> None:
        self._append_text(unescape(data))


def parse_rich_text_paragraphs(raw_html: str) -> list[list[dict[str, Any]]]:
    parser = RichTextParser()
    parser.feed(raw_html or "")
    paragraphs = [p for p in parser.paragraphs if any((r.get("text") or "").strip() for r in p)]
    return paragraphs


def paragraphs_to_reportlab_markup(paragraph_runs: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for run in paragraph_runs:
        text = run.get("text", "")
        if not text:
            continue
        text = escape(text).replace("\n", "<br/>")
        if run.get("underline"):
            text = f"<u>{text}</u>"
        if run.get("italic"):
            text = f"<i>{text}</i>"
        if run.get("bold"):
            text = f"<b>{text}</b>"
        parts.append(text)
    return "".join(parts)


def render_html_template(template_html: str, values: dict[str, Any]) -> str:
    raw = (template_html or "").strip()
    if not raw:
        return ""
    return safe_format_text(raw, values)


def normalize_inline_template_html(raw_html: str) -> str:
    html = (raw_html or "").strip()
    if not html:
        return ""
    # Rich editors often wrap inline text into p/div blocks.
    # Convert block boundaries to single line breaks and keep the rest inline.
    html = re.sub(r"(?i)<\s*p\b[^>]*>", "", html)
    html = re.sub(r"(?i)<\s*div\b[^>]*>", "", html)
    html = re.sub(r"(?i)<\s*/\s*p\s*>", "<br>", html)
    html = re.sub(r"(?i)<\s*/\s*div\s*>", "<br>", html)
    html = re.sub(r"(?i)(?:<\s*br\s*/?>\s*){2,}", "<br>", html)
    html = re.sub(r"(?i)^(?:\s*<\s*br\s*/?>\s*)+", "", html)
    html = re.sub(r"(?i)(?:\s*<\s*br\s*/?>\s*)+$", "", html)
    return html.strip()


def make_export_filename(student_name: str, semester: str, extension: str) -> str:
    base = f"Lernbrief_{student_name}_{semester}"
    safe = re.sub(r"[^A-Za-z0-9._-]", "_", base)
    return f"{safe}.{extension}"


def render_letter_pdf(letter: sqlite3.Row) -> BytesIO:
    from reportlab.lib.pagesizes import A4  # type: ignore[import-not-found]
    from reportlab.lib.styles import ParagraphStyle  # type: ignore[import-not-found]
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-not-found]

    buffer = BytesIO()
    rl_font, _ = map_export_font(str(letter["body_font_family"]))
    font_size = int(letter["body_font_size"] or 16)
    font_size = min(max(font_size, 4), 28)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=45,
        bottomMargin=45,
    )
    style = ParagraphStyle(
        "letter",
        fontName=rl_font,
        fontSize=font_size,
        leading=font_size * 1.4,
    )

    story: list[Any] = []
    paragraphs = parse_rich_text_paragraphs(str(letter["content"]))
    if not paragraphs:
        paragraphs = [[{"text": html_to_plain_text(str(letter["content"])), "bold": False, "italic": False, "underline": False}]]

    for p_runs in paragraphs:
        markup = paragraphs_to_reportlab_markup(p_runs)
        if not markup.strip():
            continue
        story.append(Paragraph(markup, style))
        story.append(Spacer(1, font_size * 0.45))

    doc.build(story)
    buffer.seek(0)
    return buffer


def render_letter_docx(letter: sqlite3.Row) -> BytesIO:
    from docx import Document  # type: ignore[import-not-found]
    from docx.shared import Pt  # type: ignore[import-not-found]

    document = Document()

    _, docx_font = map_export_font(str(letter["body_font_family"]))
    docx_size = int(letter["body_font_size"] or 16)
    docx_size = min(max(docx_size, 4), 28)

    paragraphs = parse_rich_text_paragraphs(str(letter["content"]))
    if not paragraphs:
        paragraphs = [[{"text": html_to_plain_text(str(letter["content"])), "bold": False, "italic": False, "underline": False}]]

    for p_runs in paragraphs:
        paragraph = document.add_paragraph()
        for run_data in p_runs:
            text = run_data.get("text", "")
            if text == "":
                continue
            segments = text.split("\n")
            for idx, segment in enumerate(segments):
                run = paragraph.add_run(segment)
                run.bold = bool(run_data.get("bold"))
                run.italic = bool(run_data.get("italic"))
                run.underline = bool(run_data.get("underline"))
                run.font.name = str(run_data.get("font") or docx_font)
                run.font.size = Pt(float(run_data.get("size") or docx_size))
                if idx < len(segments) - 1:
                    run.add_break()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


DEFAULT_SEMESTER = current_school_semester()


def open_browser_when_ready(url: str, host: str, port: int, timeout_seconds: int = 20) -> None:
    """Wait until the local server is reachable, then open the system browser."""
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        try:
            with socket.create_connection((host, port), timeout=0.5):
                webbrowser.open(url, new=2)
                return
        except OSError:
            time.sleep(0.3)


def create_app() -> Flask:
    app = Flask(
        __name__,
        template_folder=str(RESOURCE_DIR / "templates"),
        static_folder=str(RESOURCE_DIR / "static"),
    )
    app.config["SECRET_KEY"] = "lernbrief-hub-local-dev"

    def get_db() -> sqlite3.Connection:
        if "db" not in g:
            g.db = sqlite3.connect(DB_PATH)
            g.db.row_factory = sqlite3.Row
        return g.db

    @app.teardown_appcontext
    def close_db(_: Any) -> None:
        db = g.pop("db", None)
        if db is not None:
            db.close()

    def init_db() -> None:
        db = get_db()
        db.executescript(
            """
            CREATE TABLE IF NOT EXISTS groups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE
            );

            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id INTEGER NOT NULL,
                full_name TEXT NOT NULL,
                FOREIGN KEY (group_id) REFERENCES groups(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS competencies (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                description TEXT DEFAULT '',
                sort_order INTEGER NOT NULL DEFAULT 0
            );

            CREATE TABLE IF NOT EXISTS ratings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER NOT NULL,
                competency_id INTEGER NOT NULL,
                semester TEXT NOT NULL,
                grade INTEGER NOT NULL,
                note TEXT DEFAULT '',
                UNIQUE(student_id, competency_id, semester),
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(competency_id) REFERENCES competencies(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS sentence_templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                competency_id INTEGER NOT NULL,
                grade INTEGER NOT NULL,
                semester TEXT NOT NULL DEFAULT '*',
                sentence TEXT NOT NULL,
                FOREIGN KEY(competency_id) REFERENCES competencies(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS letters (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER NOT NULL,
                semester TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS student_semester_goals (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER NOT NULL,
                semester TEXT NOT NULL,
                goal_text TEXT NOT NULL,
                UNIQUE(student_id, semester),
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS group_semester_intros (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id INTEGER NOT NULL,
                semester TEXT NOT NULL,
                intro_text TEXT NOT NULL,
                UNIQUE(group_id, semester),
                FOREIGN KEY(group_id) REFERENCES groups(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS app_settings (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS letter_templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                header_html TEXT NOT NULL,
                footer_html TEXT NOT NULL,
                include_average_sentence INTEGER NOT NULL DEFAULT 1,
                average_sentence_template TEXT NOT NULL,
                header_position TEXT NOT NULL DEFAULT 'top',
                footer_position TEXT NOT NULL DEFAULT 'bottom',
                body_font_family TEXT NOT NULL DEFAULT 'Georgia',
                body_font_size INTEGER NOT NULL DEFAULT 16,
                is_active INTEGER NOT NULL DEFAULT 0
            );
            """
        )
        db.commit()

        letter_cols = {row["name"] for row in db.execute("PRAGMA table_info(letters)")}
        if "template_name" not in letter_cols:
            db.execute("ALTER TABLE letters ADD COLUMN template_name TEXT NOT NULL DEFAULT 'Standard'")
        if "body_font_family" not in letter_cols:
            db.execute("ALTER TABLE letters ADD COLUMN body_font_family TEXT NOT NULL DEFAULT 'Georgia'")
        if "body_font_size" not in letter_cols:
            db.execute("ALTER TABLE letters ADD COLUMN body_font_size INTEGER NOT NULL DEFAULT 16")
        db.commit()

        sentence_cols = [row["name"] for row in db.execute("PRAGMA table_info(sentence_templates)")]
        if "semester" not in sentence_cols:
            db.executescript(
                """
                ALTER TABLE sentence_templates RENAME TO sentence_templates_old;

                CREATE TABLE sentence_templates (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    competency_id INTEGER NOT NULL,
                    grade INTEGER NOT NULL,
                    semester TEXT NOT NULL DEFAULT '*',
                    sentence TEXT NOT NULL,
                    FOREIGN KEY(competency_id) REFERENCES competencies(id) ON DELETE CASCADE
                );

                INSERT INTO sentence_templates (competency_id, grade, semester, sentence)
                SELECT competency_id, grade, '*', sentence
                FROM sentence_templates_old;

                DROP TABLE sentence_templates_old;
                """
            )
            db.commit()

        competency_count = db.execute("SELECT COUNT(*) AS c FROM competencies").fetchone()["c"]
        if competency_count == 0:
            default_competencies = [
                ("Fachwissen", "Beherrscht die fachlichen Grundlagen", 1),
                ("Mitarbeit", "Bringt sich aktiv in den Unterricht ein", 2),
                ("Sozialverhalten", "Arbeitet respektvoll und kooperativ", 3),
                ("Selbstorganisation", "Plant und erledigt Aufgaben eigenständig", 4),
            ]
            db.executemany(
                "INSERT INTO competencies (name, description, sort_order) VALUES (?, ?, ?)",
                default_competencies,
            )
            db.commit()

            comp_rows = db.execute("SELECT id, name FROM competencies").fetchall()
            for comp in comp_rows:
                for grade in GRADE_OPTIONS:
                    sentence = f"In {comp['name']} erreicht { '{name}' } aktuell die Note {grade}."
                    db.execute(
                        "INSERT INTO sentence_templates (competency_id, grade, semester, sentence) VALUES (?, ?, ?, ?)",
                        (comp["id"], grade, "*", sentence),
                    )
            db.commit()

        db.execute(
            "INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)",
            ("letter_include_average_sentence", "1"),
        )
        db.execute(
            "INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)",
            (
                "letter_average_sentence_template",
                "Zusammenfassend ergibt sich eine Durchschnittsnote von {avg_grade} und damit ein insgesamt {avg_text}er Leistungsstand.",
            ),
        )
        db.execute(
            "INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)",
            (
                "letter_header_template",
                "Lernbrief für {name}\nLerngruppe: {group_name}\nHalbjahr: {semester}\n\n{full_name} hat im aktuellen Halbjahr in den vereinbarten Kompetenzbereichen insgesamt {avg_text}e Leistungen gezeigt.\n\nIm Einzelnen zeigt sich folgende Entwicklung:",
            ),
        )
        db.execute(
            "INSERT OR IGNORE INTO app_settings (key, value) VALUES (?, ?)",
            ("archived_semesters", "[]"),
        )

        template_count = db.execute("SELECT COUNT(*) AS c FROM letter_templates").fetchone()["c"]
        if template_count == 0:
            default_header = db.execute(
                "SELECT value FROM app_settings WHERE key = 'letter_header_template'"
            ).fetchone()["value"]
            default_average = db.execute(
                "SELECT value FROM app_settings WHERE key = 'letter_average_sentence_template'"
            ).fetchone()["value"]
            default_include = int(
                db.execute(
                    "SELECT value FROM app_settings WHERE key = 'letter_include_average_sentence'"
                ).fetchone()["value"]
            )
            db.execute(
                """
                INSERT INTO letter_templates (
                    name,
                    header_html,
                    footer_html,
                    include_average_sentence,
                    average_sentence_template,
                    header_position,
                    footer_position,
                    body_font_family,
                    body_font_size,
                    is_active
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
                """,
                (
                    "Standard",
                    default_header.replace("\n", "<br>"),
                    "Für das kommende Halbjahr werden wir den eingeschlagenen Entwicklungsweg mit {name} kontinuierlich fortsetzen.",
                    default_include,
                    default_average,
                    "top",
                    "bottom",
                    "Georgia",
                    16,
                ),
            )

        # Migration: Make the lead-in sentence editable in the header for existing templates.
        default_header_lead = (
            "{full_name} hat im aktuellen Halbjahr in den vereinbarten Kompetenzbereichen "
            "insgesamt {avg_text}e Leistungen gezeigt.<br><br>Im Einzelnen zeigt sich folgende Entwicklung:"
        )
        existing_templates = db.execute(
            "SELECT id, header_html FROM letter_templates ORDER BY id ASC"
        ).fetchall()
        for tpl_row in existing_templates:
            current_header = (tpl_row["header_html"] or "").strip()
            if "vereinbarten Kompetenzbereichen" in current_header:
                continue
            if current_header:
                updated_header = f"{current_header}<br><br>{default_header_lead}"
            else:
                updated_header = (
                    "Lernbrief für {name}<br>Lerngruppe: {group_name}<br>Halbjahr: {semester}<br><br>"
                    f"{default_header_lead}"
                )
            db.execute(
                "UPDATE letter_templates SET header_html = ? WHERE id = ?",
                (updated_header, tpl_row["id"]),
            )

        db.execute("UPDATE letter_templates SET is_active = 0 WHERE id NOT IN (SELECT id FROM letter_templates WHERE is_active = 1)")
        active_count = db.execute("SELECT COUNT(*) AS c FROM letter_templates WHERE is_active = 1").fetchone()["c"]
        if active_count == 0:
            db.execute("UPDATE letter_templates SET is_active = 1 WHERE id = (SELECT id FROM letter_templates ORDER BY id ASC LIMIT 1)")
        db.commit()

    @app.before_request
    def ensure_db() -> None:
        init_db()

    def query_groups() -> list[sqlite3.Row]:
        return get_db().execute(
            """
            SELECT g.id, g.name, COUNT(s.id) AS student_count
            FROM groups g
            LEFT JOIN students s ON s.group_id = g.id
            GROUP BY g.id, g.name
            ORDER BY g.name ASC
            """
        ).fetchall()

    def query_competencies() -> list[sqlite3.Row]:
        return get_db().execute(
            "SELECT * FROM competencies ORDER BY sort_order ASC, name ASC"
        ).fetchall()

    def query_overview_data() -> dict[str, Any]:
        db = get_db()
        summary = db.execute(
            """
            SELECT
                (SELECT COUNT(*) FROM groups) AS group_count,
                (SELECT COUNT(*) FROM students) AS student_count,
                (SELECT COUNT(*) FROM competencies) AS competency_count,
                (SELECT COUNT(*) FROM ratings) AS rating_count,
                (SELECT COUNT(*) FROM letters) AS letter_count,
                (SELECT COUNT(DISTINCT student_id) FROM ratings) AS rated_student_count,
                (SELECT COUNT(DISTINCT student_id) FROM letters) AS letter_student_count,
                (SELECT COUNT(DISTINCT semester) FROM ratings) AS rating_semester_count,
                (SELECT COUNT(DISTINCT semester) FROM letters) AS letter_semester_count,
                (SELECT COUNT(*) FROM ratings WHERE semester = ?) AS current_semester_rating_count,
                (SELECT COUNT(*) FROM letters WHERE semester = ?) AS current_semester_letter_count
            """,
            (DEFAULT_SEMESTER, DEFAULT_SEMESTER),
        ).fetchone()

        largest_groups = db.execute(
            """
            SELECT g.id, g.name, COUNT(s.id) AS student_count
            FROM groups g
            LEFT JOIN students s ON s.group_id = g.id
            GROUP BY g.id, g.name
            ORDER BY student_count DESC, g.name ASC
            LIMIT 5
            """
        ).fetchall()

        semester_activity = db.execute(
            """
            SELECT semester, SUM(rating_count) AS rating_count, SUM(letter_count) AS letter_count
            FROM (
                SELECT semester, COUNT(*) AS rating_count, 0 AS letter_count
                FROM ratings
                GROUP BY semester

                UNION ALL

                SELECT semester, 0 AS rating_count, COUNT(*) AS letter_count
                FROM letters
                GROUP BY semester
            ) activity
            GROUP BY semester
            ORDER BY semester DESC
            LIMIT 6
            """
        ).fetchall()

        recent_letters = db.execute(
            """
            SELECT l.id, l.semester, l.created_at, s.full_name
            FROM letters l
            JOIN students s ON s.id = l.student_id
            ORDER BY l.created_at DESC
            LIMIT 5
            """
        ).fetchall()

        return {
            "summary": summary,
            "largest_groups": largest_groups,
            "semester_activity": semester_activity,
            "recent_letters": recent_letters,
            "default_semester": DEFAULT_SEMESTER,
        }

    def get_setting(key: str, default: str = "") -> str:
        row = get_db().execute("SELECT value FROM app_settings WHERE key = ?", (key,)).fetchone()
        if row is None:
            return default
        return row["value"]

    def set_setting(key: str, value: str) -> None:
        get_db().execute(
            """
            INSERT INTO app_settings (key, value)
            VALUES (?, ?)
            ON CONFLICT(key) DO UPDATE SET value = excluded.value
            """,
            (key, value),
        )

    def get_archived_semesters() -> set[str]:
        raw = get_setting("archived_semesters", "[]")
        try:
            loaded = json.loads(raw)
        except json.JSONDecodeError:
            return set()
        if not isinstance(loaded, list):
            return set()
        return {str(item) for item in loaded if parse_semester(str(item))}

    def set_archived_semesters(values: set[str]) -> None:
        serialized = json.dumps(sorted(values, key=semester_sort_key, reverse=True))
        set_setting("archived_semesters", serialized)

    def get_active_letter_template() -> sqlite3.Row:
        db = get_db()
        tpl = db.execute(
            "SELECT * FROM letter_templates WHERE is_active = 1 ORDER BY id ASC LIMIT 1"
        ).fetchone()
        if tpl is None:
            tpl = db.execute("SELECT * FROM letter_templates ORDER BY id ASC LIMIT 1").fetchone()
        return tpl

    def build_letter(student_id: int, semester: str) -> str:
        db = get_db()
        student = db.execute(
            """
            SELECT s.id, s.group_id, s.full_name, g.name AS group_name
            FROM students s
            JOIN groups g ON g.id = s.group_id
            WHERE s.id = ?
            """,
            (student_id,),
        ).fetchone()

        ratings = db.execute(
            """
            SELECT c.name AS competency_name, r.competency_id, r.grade, r.note
            FROM ratings r
            JOIN competencies c ON c.id = r.competency_id
            WHERE r.student_id = ? AND r.semester = ?
            ORDER BY c.sort_order ASC, c.name ASC
            """,
            (student_id, semester),
        ).fetchall()

        if not ratings:
            raise ValueError("Keine Bewertungen für dieses Halbjahr vorhanden.")

        avg_grade = round(sum(row["grade"] for row in ratings) / len(ratings), 2)
        avg_text = (
            "sehr gut" if avg_grade <= 1.5 else
            "gut" if avg_grade <= 2.5 else
            "befriedigend" if avg_grade <= 3.5 else
            "ausreichend" if avg_grade <= 4.5 else
            "verbesserungsbedürftig"
        )

        intro_row = db.execute(
            """
            SELECT intro_text
            FROM group_semester_intros
            WHERE group_id = ? AND semester = ?
            """,
            (student["group_id"], semester),
        ).fetchone()
        semester_intro = ""
        if intro_row is not None:
            semester_intro = " ".join(part.strip() for part in intro_row["intro_text"].splitlines() if part.strip())

        goals_row = db.execute(
            """
            SELECT goal_text
            FROM student_semester_goals
            WHERE student_id = ? AND semester = ?
            """,
            (student_id, semester),
        ).fetchone()
        semester_goals = ""
        if goals_row is not None:
            semester_goals = " ".join(part.strip() for part in goals_row["goal_text"].splitlines() if part.strip())

        tpl = get_active_letter_template()
        format_values = {
            "name": student["full_name"],
            "full_name": student["full_name"],
            "group_name": student["group_name"],
            "semester": semester,
            "avg_grade": avg_grade,
            "avg_text": avg_text,
        }

        header_html = normalize_inline_template_html(
            render_html_template(tpl["header_html"], format_values)
        )
        if not header_html.strip():
            fallback = (
                "Lernbrief für {name}<br>Lerngruppe: {group_name}<br>Halbjahr: {semester}<br><br>"
                "{full_name} hat im aktuellen Halbjahr in den vereinbarten Kompetenzbereichen insgesamt {avg_text}e Leistungen gezeigt.<br><br>"
                "Im Einzelnen zeigt sich folgende Entwicklung:"
            )
            header_html = normalize_inline_template_html(render_html_template(fallback, format_values))

        intro_html = ""
        if semester_intro:
            intro_html = ensure_sentence_punctuation(semester_intro)

        body_paragraphs_html: list[str] = []
        for idx, row in enumerate(ratings, start=1):
            template_rows = db.execute(
                """
                SELECT sentence
                FROM sentence_templates
                WHERE competency_id = ?
                  AND grade = ?
                ORDER BY id ASC
                """,
                (row["competency_id"], row["grade"]),
            ).fetchall()
            if template_rows:
                template = random.choice(template_rows)["sentence"]
            else:
                template = f"In {row['competency_name']} liegt die Leistung bei der Note {row['grade']}."
            sentence = template.replace("{name}", student["full_name"])
            if row["note"]:
                sentence = f"{sentence} Hinweis: {row['note']}"

            # Normalize multi-line template blocks into readable report paragraphs.
            paragraph = " ".join(part.strip() for part in sentence.splitlines() if part.strip())
            paragraph = ensure_sentence_punctuation(paragraph)
            body_paragraphs_html.append(paragraph)

            # Add a small paragraph break every two competency blocks.
            if idx % 2 == 0 and idx < len(ratings):
                body_paragraphs_html.append("")

        footer_parts: list[str] = []
        if int(tpl["include_average_sentence"]) == 1:
            average_sentence = safe_format_text(tpl["average_sentence_template"], format_values)
            footer_parts.append(ensure_sentence_punctuation(average_sentence))

        custom_footer = normalize_inline_template_html(
            render_html_template(tpl["footer_html"], format_values)
        )
        if custom_footer.strip():
            footer_parts.append(custom_footer)
        else:
            footer_parts.append(
                f"Für das kommende Halbjahr werden wir den eingeschlagenen Entwicklungsweg mit {student['full_name']} kontinuierlich fortsetzen."
            )

        if semester_goals:
            footer_parts.append(f"Halbjahresziele:<br>{ensure_sentence_punctuation(semester_goals)}")

        header_position = tpl["header_position"] if tpl["header_position"] in {"top", "after_intro"} else "top"
        footer_position = tpl["footer_position"] if tpl["footer_position"] in {"bottom", "after_header"} else "bottom"
        body_font_family = tpl["body_font_family"] or "Georgia"
        body_font_size = int(tpl["body_font_size"] or 16)
        body_font_size = min(max(body_font_size, 4), 28)

        html_parts: list[str] = []

        if header_position == "top" and header_html:
            html_parts.append(f"<div class='letter-header'>{header_html}</div>")

        if footer_position == "after_header":
            for part in footer_parts:
                if part.strip():
                    html_parts.append(f"<div class='letter-footer'>{part}</div>")

        if intro_html:
            html_parts.append(f"<p>{intro_html}</p>")

        if header_position == "after_intro" and header_html:
            html_parts.append(f"<div class='letter-header'>{header_html}</div>")

        body_clean = [p for p in body_paragraphs_html if p.strip()]
        for paragraph in body_clean:
            html_parts.append(f"<p>{paragraph}</p>")

        if footer_position == "bottom":
            for part in footer_parts:
                if part.strip():
                    html_parts.append(f"<div class='letter-footer'>{part}</div>")

        content_html = "".join(html_parts)
        wrapped = (
            f"<div style='font-family:{body_font_family};font-size:{body_font_size}px;line-height:1.55;'>"
            f"{content_html}"
            "</div>"
        )
        return wrapped

    @app.route("/")
    def index() -> str:
        search_query = request.args.get("q", "").strip()
        groups = query_groups()
        recent_letters = get_db().execute(
            """
            SELECT l.id, l.semester, l.created_at, s.full_name
            FROM letters l
            JOIN students s ON s.id = l.student_id
            ORDER BY l.created_at DESC
            LIMIT 10
            """
        ).fetchall()

        student_results: list[sqlite3.Row] = []
        if search_query:
            student_results = get_db().execute(
                """
                SELECT s.id AS student_id, s.full_name, g.id AS group_id, g.name AS group_name
                FROM students s
                JOIN groups g ON g.id = s.group_id
                WHERE s.full_name LIKE ? OR g.name LIKE ?
                ORDER BY s.full_name ASC
                LIMIT 50
                """,
                (f"%{search_query}%", f"%{search_query}%"),
            ).fetchall()

        return render_template(
            "index.html",
            groups=groups,
            recent_letters=recent_letters,
            search_query=search_query,
            student_results=student_results,
            default_semester=DEFAULT_SEMESTER,
        )

    @app.route("/overview")
    def overview() -> str:
        return render_template("overview.html", **query_overview_data())

    @app.route("/data")
    def data_management() -> str:
        db = get_db()
        archived_semesters = get_archived_semesters()

        semester_rows = db.execute(
            """
            SELECT sem.semester,
                   COALESCE(r.rating_count, 0) AS rating_count,
                   COALESCE(l.letter_count, 0) AS letter_count
            FROM (
                SELECT semester FROM ratings
                UNION
                SELECT semester FROM letters
            ) sem
            LEFT JOIN (
                SELECT semester, COUNT(*) AS rating_count
                FROM ratings
                GROUP BY semester
            ) r ON r.semester = sem.semester
            LEFT JOIN (
                SELECT semester, COUNT(*) AS letter_count
                FROM letters
                GROUP BY semester
            ) l ON l.semester = sem.semester
            ORDER BY sem.semester DESC
            """
        ).fetchall()

        semester_items = [
            {
                "semester": row["semester"],
                "rating_count": row["rating_count"],
                "letter_count": row["letter_count"],
                "archived": row["semester"] in archived_semesters,
            }
            for row in semester_rows
        ]

        db_size_kb = 0
        if DB_PATH.exists():
            db_size_kb = int(DB_PATH.stat().st_size / 1024)

        return render_template(
            "data_management.html",
            semester_items=semester_items,
            archived_count=len(archived_semesters),
            db_file_name=DB_PATH.name,
            db_size_kb=db_size_kb,
        )

    @app.route("/data/archive-toggle", methods=["POST"])
    def toggle_archive_semester() -> Any:
        semester = (request.form.get("semester") or "").strip()
        if not parse_semester(semester):
            flash("Ungültiges Halbjahr.", "error")
            return redirect(url_for("data_management"))

        action = request.form.get("action", "archive")
        archived_semesters = get_archived_semesters()

        if action == "archive":
            archived_semesters.add(semester)
            flash(f"Halbjahr {semester} wurde archiviert.", "success")
        else:
            archived_semesters.discard(semester)
            flash(f"Halbjahr {semester} wurde reaktiviert.", "success")

        set_archived_semesters(archived_semesters)
        get_db().commit()
        return redirect(url_for("data_management"))

    @app.route("/data/backup")
    def download_backup() -> Any:
        if not DB_PATH.exists():
            flash("Keine Datenbank für Backup gefunden.", "error")
            return redirect(url_for("data_management"))

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"lernbrief_hub_backup_{timestamp}.db"
        return send_file(DB_PATH, as_attachment=True, download_name=file_name, mimetype="application/octet-stream")

    @app.route("/data/restore", methods=["POST"])
    def restore_backup() -> Any:
        upload = request.files.get("backup_file")
        if upload is None or not upload.filename:
            flash("Bitte eine Backup-Datei auswählen.", "error")
            return redirect(url_for("data_management"))

        restore_temp_path = DATA_DIR / f"lernbrief_hub_restore_{int(time.time())}.tmp"
        try:
            upload.save(restore_temp_path)

            test_conn = sqlite3.connect(restore_temp_path)
            try:
                existing_tables = {
                    row[0]
                    for row in test_conn.execute(
                        "SELECT name FROM sqlite_master WHERE type='table'"
                    ).fetchall()
                }
            finally:
                test_conn.close()

            required_tables = {
                "groups",
                "students",
                "competencies",
                "ratings",
                "sentence_templates",
                "letters",
                "app_settings",
            }
            missing_tables = required_tables - existing_tables
            if missing_tables:
                flash("Backup-Datei ist nicht kompatibel (fehlende Tabellen).", "error")
                return redirect(url_for("data_management"))

            open_db = g.pop("db", None)
            if open_db is not None:
                open_db.close()

            if DB_PATH.exists():
                safety_copy = DATA_DIR / f"lernbrief_hub_pre_restore_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bak"
                shutil.copy2(DB_PATH, safety_copy)

            os.replace(restore_temp_path, DB_PATH)
            flash("Backup wurde erfolgreich wiederhergestellt.", "success")
            return redirect(url_for("data_management"))
        except Exception:
            flash("Wiederherstellung fehlgeschlagen. Bitte Backup-Datei prüfen.", "error")
            return redirect(url_for("data_management"))
        finally:
            if restore_temp_path.exists():
                restore_temp_path.unlink()

    @app.route("/groups/create", methods=["POST"])
    def create_group() -> Any:
        name = request.form.get("name", "").strip()
        if not name:
            flash("Bitte einen Gruppennamen eingeben.", "error")
            return redirect(url_for("index"))

        try:
            get_db().execute("INSERT INTO groups (name) VALUES (?)", (name,))
            get_db().commit()
            flash("Lerngruppe erstellt.", "success")
        except sqlite3.IntegrityError:
            flash("Diese Lerngruppe existiert bereits.", "error")
        return redirect(url_for("index"))

    @app.route("/groups/new")
    def new_group() -> str:
        return render_template("group_new.html")

    @app.route("/groups/<int:group_id>")
    def group_detail(group_id: int) -> str:
        db = get_db()
        group = db.execute("SELECT * FROM groups WHERE id = ?", (group_id,)).fetchone()
        if group is None:
            flash("Lerngruppe nicht gefunden.", "error")
            return redirect(url_for("index"))

        students = db.execute(
            "SELECT * FROM students WHERE group_id = ? ORDER BY full_name ASC", (group_id,)
        ).fetchall()

        semester = normalize_semester(request.args.get("semester", DEFAULT_SEMESTER))
        semester_options = school_semester_options(db=db)
        if semester not in semester_options:
            semester_options.append(semester)
            semester_options = sorted(semester_options, key=semester_sort_key, reverse=True)

        intro_row = db.execute(
            """
            SELECT intro_text
            FROM group_semester_intros
            WHERE group_id = ? AND semester = ?
            """,
            (group_id, semester),
        ).fetchone()
        semester_intro_text = intro_row["intro_text"] if intro_row is not None else ""

        return render_template(
            "group_detail.html",
            group=group,
            students=students,
            default_semester=DEFAULT_SEMESTER,
            semester=semester,
            semester_options=semester_options,
            semester_intro_text=semester_intro_text,
        )

    @app.route("/groups/<int:group_id>/semester-text", methods=["POST"])
    def save_group_semester_text(group_id: int) -> Any:
        db = get_db()
        group = db.execute("SELECT id FROM groups WHERE id = ?", (group_id,)).fetchone()
        if group is None:
            flash("Lerngruppe nicht gefunden.", "error")
            return redirect(url_for("index"))

        semester = normalize_semester(request.form.get("semester", DEFAULT_SEMESTER))
        if semester in get_archived_semesters():
            flash("Archivierte Halbjahre sind schreibgeschützt. Bitte Halbjahr reaktivieren.", "error")
            return redirect(url_for("group_detail", group_id=group_id, semester=semester))

        intro_text = request.form.get("semester_intro_text", "").strip()
        if intro_text:
            db.execute(
                """
                INSERT INTO group_semester_intros (group_id, semester, intro_text)
                VALUES (?, ?, ?)
                ON CONFLICT(group_id, semester)
                DO UPDATE SET intro_text = excluded.intro_text
                """,
                (group_id, semester, intro_text),
            )
        else:
            db.execute(
                "DELETE FROM group_semester_intros WHERE group_id = ? AND semester = ?",
                (group_id, semester),
            )

        db.commit()
        flash("Halbjahrestext der Lerngruppe gespeichert.", "success")
        return redirect(url_for("group_detail", group_id=group_id, semester=semester))

    @app.route("/groups/<int:group_id>/students/create", methods=["POST"])
    def create_student(group_id: int) -> Any:
        full_name = request.form.get("full_name", "").strip()
        if not full_name:
            flash("Bitte einen Schülernamen eingeben.", "error")
            return redirect(url_for("group_detail", group_id=group_id))

        get_db().execute(
            "INSERT INTO students (group_id, full_name) VALUES (?, ?)",
            (group_id, full_name),
        )
        get_db().commit()
        flash("Schüler hinzugefügt.", "success")
        return redirect(url_for("group_detail", group_id=group_id))

    @app.route("/competencies", methods=["GET", "POST"])
    def competencies() -> Any:
        db = get_db()

        if request.method == "POST":
            action = request.form.get("action", "create")

            if action == "update":
                competency_id = int(request.form.get("competency_id", "0") or 0)
                name = request.form.get("name", "").strip()
                description = request.form.get("description", "").strip()
                sort_order = int(request.form.get("sort_order", "0") or 0)

                if not competency_id or not name:
                    flash("Bitte gültige Kompetenzdaten angeben.", "error")
                    return redirect(url_for("competencies"))

                try:
                    db.execute(
                        """
                        UPDATE competencies
                        SET name = ?, description = ?, sort_order = ?
                        WHERE id = ?
                        """,
                        (name, description, sort_order, competency_id),
                    )
                    db.commit()
                    flash("Kompetenz aktualisiert.", "success")
                except sqlite3.IntegrityError:
                    flash("Eine Kompetenz mit diesem Namen existiert bereits.", "error")

                return redirect(url_for("competencies"))

            name = request.form.get("name", "").strip()
            description = request.form.get("description", "").strip()
            sort_order = int(request.form.get("sort_order", "0") or 0)

            if not name:
                flash("Kompetenzname fehlt.", "error")
                return redirect(url_for("competencies"))

            try:
                cur = db.execute(
                    "INSERT INTO competencies (name, description, sort_order) VALUES (?, ?, ?)",
                    (name, description, sort_order),
                )
                competency_id = cur.lastrowid
                for grade in GRADE_OPTIONS:
                    db.execute(
                        "INSERT INTO sentence_templates (competency_id, grade, semester, sentence) VALUES (?, ?, ?, ?)",
                        (competency_id, grade, "*", f"In {name} erreicht {{name}} aktuell die Note {grade}."),
                    )
                db.commit()
                flash("Kompetenz erstellt.", "success")
            except sqlite3.IntegrityError:
                flash("Diese Kompetenz existiert bereits.", "error")

            return redirect(url_for("competencies"))

        rows = query_competencies()
        return render_template("competencies.html", competencies=rows)

    @app.route("/templates", methods=["GET", "POST"])
    def templates_editor() -> Any:
        db = get_db()

        if request.method == "POST":
            action = request.form.get("action", "update")

            if action == "create":
                competency_id = int(request.form.get("competency_id", "0") or 0)
                grade = int(request.form.get("grade", "0") or 0)
                sentence = request.form.get("sentence", "").strip()

                if not competency_id or grade not in GRADE_OPTIONS or not sentence:
                    flash("Bitte Kompetenz, Note und Satz angeben.", "error")
                    return redirect(url_for("templates_editor"))

                db.execute(
                    "INSERT INTO sentence_templates (competency_id, grade, semester, sentence) VALUES (?, ?, ?, ?)",
                    (competency_id, grade, "*", sentence),
                )
                db.commit()
                flash("Satzbaustein hinzugefügt.", "success")
                return redirect(url_for("templates_editor"))

            if action == "delete":
                template_id = int(request.form["template_id"])
                db.execute("DELETE FROM sentence_templates WHERE id = ?", (template_id,))
                db.commit()
                flash("Satzbaustein gelöscht.", "success")
                return redirect(url_for("templates_editor"))

            template_id = int(request.form["template_id"])
            sentence = request.form.get("sentence", "").strip()

            if not sentence:
                flash("Satz darf nicht leer sein.", "error")
                return redirect(url_for("templates_editor"))

            db.execute(
                "UPDATE sentence_templates SET sentence = ? WHERE id = ?",
                (sentence, template_id),
            )
            db.commit()
            flash("Satzbaustein gespeichert.", "success")
            return redirect(url_for("templates_editor"))

        rows = db.execute(
            """
            SELECT st.id, st.grade, st.sentence, c.name AS competency_name, c.id AS competency_id
            FROM sentence_templates st
            JOIN competencies c ON c.id = st.competency_id
            ORDER BY c.sort_order ASC, c.name ASC, st.grade ASC
            """
        ).fetchall()
        competencies_rows = query_competencies()
        return render_template(
            "templates.html",
            templates=rows,
            competencies=competencies_rows,
            grade_options=GRADE_OPTIONS,
        )

    @app.route("/letter-templates", methods=["GET", "POST"])
    def letter_templates_editor() -> Any:
        db = get_db()

        if request.method == "POST":
            action = request.form.get("action", "save")
            selected_id = int(request.form.get("template_id", "0") or 0)

            if action == "create":
                name = request.form.get("new_template_name", "").strip()
                if not name:
                    flash("Bitte einen Namen für die Vorlage eingeben.", "error")
                    return redirect(url_for("letter_templates_editor"))
                try:
                    db.execute(
                        """
                        INSERT INTO letter_templates (
                            name, header_html, footer_html, include_average_sentence,
                            average_sentence_template, header_position, footer_position,
                            body_font_family, body_font_size, is_active
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 0)
                        """,
                        (
                            name,
                            "Lernbrief für {name}<br>Lerngruppe: {group_name}<br>Halbjahr: {semester}",
                            "Für das kommende Halbjahr werden wir den eingeschlagenen Entwicklungsweg mit {name} kontinuierlich fortsetzen.",
                            1,
                            "Zusammenfassend ergibt sich eine Durchschnittsnote von {avg_grade} und damit ein insgesamt {avg_text}er Leistungsstand.",
                            "top",
                            "bottom",
                            "Georgia",
                            16,
                        ),
                    )
                    db.commit()
                    flash("Neue Lernbriefvorlage erstellt.", "success")
                except sqlite3.IntegrityError:
                    flash("Eine Vorlage mit diesem Namen existiert bereits.", "error")
                return redirect(url_for("letter_templates_editor"))

            if selected_id == 0:
                flash("Bitte eine Vorlage auswählen.", "error")
                return redirect(url_for("letter_templates_editor"))

            if action == "activate":
                db.execute("UPDATE letter_templates SET is_active = 0")
                db.execute("UPDATE letter_templates SET is_active = 1 WHERE id = ?", (selected_id,))
                db.commit()
                flash("Vorlage als aktiv gesetzt.", "success")
                return redirect(url_for("letter_templates_editor", template_id=selected_id))

            if action == "delete":
                count = db.execute("SELECT COUNT(*) AS c FROM letter_templates").fetchone()["c"]
                if count <= 1:
                    flash("Mindestens eine Lernbriefvorlage muss erhalten bleiben.", "error")
                    return redirect(url_for("letter_templates_editor", template_id=selected_id))
                was_active = db.execute(
                    "SELECT is_active FROM letter_templates WHERE id = ?",
                    (selected_id,),
                ).fetchone()
                db.execute("DELETE FROM letter_templates WHERE id = ?", (selected_id,))
                if was_active and int(was_active["is_active"]) == 1:
                    db.execute(
                        "UPDATE letter_templates SET is_active = 1 WHERE id = (SELECT id FROM letter_templates ORDER BY id ASC LIMIT 1)"
                    )
                db.commit()
                flash("Vorlage gelöscht.", "success")
                return redirect(url_for("letter_templates_editor"))

            name = request.form.get("name", "").strip()
            header_html = request.form.get("header_html", "").strip()
            footer_html = request.form.get("footer_html", "").strip()
            include_average_sentence = 1 if request.form.get("include_average_sentence") == "on" else 0
            average_sentence_template = request.form.get("average_sentence_template", "").strip()
            header_position = request.form.get("header_position", "top")
            footer_position = request.form.get("footer_position", "bottom")
            body_font_family = request.form.get("body_font_family", "Georgia").strip() or "Georgia"
            body_font_size = int(request.form.get("body_font_size", "16") or 16)

            if not name:
                flash("Der Vorlagenname darf nicht leer sein.", "error")
                return redirect(url_for("letter_templates_editor", template_id=selected_id))
            if not header_html:
                header_html = "Lernbrief für {name}<br>Lerngruppe: {group_name}<br>Halbjahr: {semester}"
            if not footer_html:
                footer_html = "Für das kommende Halbjahr werden wir den eingeschlagenen Entwicklungsweg mit {name} kontinuierlich fortsetzen."
            if not average_sentence_template:
                average_sentence_template = (
                    "Zusammenfassend ergibt sich eine Durchschnittsnote von {avg_grade} "
                    "und damit ein insgesamt {avg_text}er Leistungsstand."
                )
            if header_position not in {"top", "after_intro"}:
                header_position = "top"
            if footer_position not in {"bottom", "after_header"}:
                footer_position = "bottom"
            body_font_size = min(max(body_font_size, 4), 28)

            try:
                db.execute(
                    """
                    UPDATE letter_templates
                    SET name = ?,
                        header_html = ?,
                        footer_html = ?,
                        include_average_sentence = ?,
                        average_sentence_template = ?,
                        header_position = ?,
                        footer_position = ?,
                        body_font_family = ?,
                        body_font_size = ?
                    WHERE id = ?
                    """,
                    (
                        name,
                        header_html,
                        footer_html,
                        include_average_sentence,
                        average_sentence_template,
                        header_position,
                        footer_position,
                        body_font_family,
                        body_font_size,
                        selected_id,
                    ),
                )
                db.commit()
                flash("Lernbriefvorlage gespeichert.", "success")
            except sqlite3.IntegrityError:
                flash("Eine Vorlage mit diesem Namen existiert bereits.", "error")

            return redirect(url_for("letter_templates_editor", template_id=selected_id))

        templates = db.execute(
            "SELECT * FROM letter_templates ORDER BY id ASC"
        ).fetchall()
        requested_id = int(request.args.get("template_id", "0") or 0)
        selected_template = None
        if requested_id:
            selected_template = next((tpl for tpl in templates if tpl["id"] == requested_id), None)
        if selected_template is None and templates:
            selected_template = next((tpl for tpl in templates if tpl["is_active"] == 1), templates[0])

        return render_template(
            "letter_templates.html",
            templates=templates,
            selected_template=selected_template,
        )

    @app.route("/students/<int:student_id>/ratings", methods=["GET", "POST"])
    def ratings(student_id: int) -> Any:
        db = get_db()
        archived_semesters = get_archived_semesters()
        semester = normalize_semester(request.values.get("semester", DEFAULT_SEMESTER))
        semester_options = school_semester_options(db=db)
        semester_options = [item for item in semester_options if item not in archived_semesters]
        if semester not in semester_options:
            semester_options.append(semester)
            semester_options = sorted(semester_options, key=semester_sort_key, reverse=True)

        student = db.execute(
            """
            SELECT s.id, s.group_id, s.full_name, g.name AS group_name
            FROM students s
            JOIN groups g ON g.id = s.group_id
            WHERE s.id = ?
            """,
            (student_id,),
        ).fetchone()
        if student is None:
            flash("Schüler nicht gefunden.", "error")
            return redirect(url_for("index"))

        competencies_rows = query_competencies()

        if request.method == "POST":
            if semester in archived_semesters:
                flash("Archivierte Halbjahre sind schreibgeschützt. Bitte Halbjahr reaktivieren.", "error")
                return redirect(url_for("ratings", student_id=student_id, semester=semester))

            action = request.form.get("action", "save_ratings")

            if action == "save_student_goal":
                goal_text = request.form.get("semester_goal_text", "").strip()

                if goal_text:
                    db.execute(
                        """
                        INSERT INTO student_semester_goals (student_id, semester, goal_text)
                        VALUES (?, ?, ?)
                        ON CONFLICT(student_id, semester)
                        DO UPDATE SET goal_text = excluded.goal_text
                        """,
                        (student_id, semester, goal_text),
                    )
                else:
                    db.execute(
                        "DELETE FROM student_semester_goals WHERE student_id = ? AND semester = ?",
                        (student_id, semester),
                    )

                db.commit()
                flash("Halbjahresziel gespeichert.", "success")
                return redirect(url_for("ratings", student_id=student_id, semester=semester))

            for comp in competencies_rows:
                grade_val = request.form.get(f"grade_{comp['id']}")
                note_val = request.form.get(f"note_{comp['id']}", "").strip()
                if not grade_val:
                    continue
                grade = int(grade_val)

                db.execute(
                    """
                    INSERT INTO ratings (student_id, competency_id, semester, grade, note)
                    VALUES (?, ?, ?, ?, ?)
                    ON CONFLICT(student_id, competency_id, semester)
                    DO UPDATE SET grade = excluded.grade, note = excluded.note
                    """,
                    (student_id, comp["id"], semester, grade, note_val),
                )
            db.commit()
            flash("Bewertungen gespeichert.", "success")
            return redirect(url_for("ratings", student_id=student_id, semester=semester))

        existing_ratings = db.execute(
            "SELECT competency_id, grade, note FROM ratings WHERE student_id = ? AND semester = ?",
            (student_id, semester),
        ).fetchall()
        ratings_map = {row["competency_id"]: row for row in existing_ratings}

        letters = db.execute(
            "SELECT id, semester, created_at FROM letters WHERE student_id = ? ORDER BY created_at DESC",
            (student_id,),
        ).fetchall()

        rating_history = db.execute(
            """
            SELECT r.semester, c.name AS competency_name, r.grade, r.note
            FROM ratings r
            JOIN competencies c ON c.id = r.competency_id
            WHERE r.student_id = ?
            ORDER BY r.semester DESC, c.sort_order ASC, c.name ASC
            """,
            (student_id,),
        ).fetchall()

        semester_stats = db.execute(
            """
            SELECT semester, ROUND(AVG(grade), 2) AS avg_grade, COUNT(*) AS rating_count
            FROM ratings
            WHERE student_id = ?
            GROUP BY semester
            ORDER BY semester DESC
            """,
            (student_id,),
        ).fetchall()

        letter_stats_rows = db.execute(
            """
            SELECT semester, COUNT(*) AS letter_count, MAX(created_at) AS last_letter_created_at
            FROM letters
            WHERE student_id = ?
            GROUP BY semester
            """,
            (student_id,),
        ).fetchall()
        letter_stats = {
            row["semester"]: {
                "letter_count": row["letter_count"],
                "last_letter_created_at": row["last_letter_created_at"],
            }
            for row in letter_stats_rows
        }

        goal_row = db.execute(
            """
            SELECT goal_text
            FROM student_semester_goals
            WHERE student_id = ? AND semester = ?
            """,
            (student_id, semester),
        ).fetchone()
        semester_goal_text = goal_row["goal_text"] if goal_row is not None else ""

        semester_overview: list[dict[str, Any]] = []
        for row in semester_stats:
            letters_info = letter_stats.get(row["semester"], {})
            semester_overview.append(
                {
                    "semester": row["semester"],
                    "avg_grade": row["avg_grade"],
                    "rating_count": row["rating_count"],
                    "letter_count": letters_info.get("letter_count", 0),
                    "last_letter_created_at": letters_info.get("last_letter_created_at", "-"),
                }
            )

        trend_rows = db.execute(
            """
            SELECT r.semester, r.competency_id, c.name AS competency_name, r.grade
            FROM ratings r
            JOIN competencies c ON c.id = r.competency_id
            WHERE r.student_id = ?
            """,
            (student_id,),
        ).fetchall()

        trend_semesters = sorted(
            {row["semester"] for row in trend_rows if parse_semester(row["semester"])},
            key=semester_sort_key,
        )

        grades_by_competency: dict[int, dict[str, int]] = {}
        for row in trend_rows:
            competency_id = row["competency_id"]
            if competency_id not in grades_by_competency:
                grades_by_competency[competency_id] = {}
            grades_by_competency[competency_id][row["semester"]] = row["grade"]

        competency_trends: list[dict[str, Any]] = []
        for comp in competencies_rows:
            grades_for_semesters: list[int | None] = [
                grades_by_competency.get(comp["id"], {}).get(sem)
                for sem in trend_semesters
            ]

            points: list[dict[str, Any]] = []
            for sem, grade in zip(trend_semesters, grades_for_semesters):
                score_percent = 0 if grade is None else round(((7 - grade) / 6) * 100, 1)
                points.append(
                    {
                        "semester": sem,
                        "grade": grade,
                        "score_percent": score_percent,
                        "grade_class": f"grade-{grade}" if grade is not None else "grade-none",
                    }
                )

            graded_values = [grade for grade in grades_for_semesters if grade is not None]
            latest_grade = graded_values[-1] if graded_values else None
            previous_grade = graded_values[-2] if len(graded_values) >= 2 else None

            delta_text = "-"
            trend_class = "trend-neutral"
            if latest_grade is not None and previous_grade is not None:
                delta = latest_grade - previous_grade
                if delta < 0:
                    delta_text = f"Verbessert um {abs(delta)}"
                    trend_class = "trend-up"
                elif delta > 0:
                    delta_text = f"Verschlechtert um {delta}"
                    trend_class = "trend-down"
                else:
                    delta_text = "Konstant"

            competency_trends.append(
                {
                    "name": comp["name"],
                    "points": points,
                    "latest_grade": latest_grade,
                    "delta_text": delta_text,
                    "trend_class": trend_class,
                }
            )

        return render_template(
            "ratings.html",
            student=student,
            semester=semester,
            semester_options=semester_options,
            competencies=competencies_rows,
            ratings_map=ratings_map,
            grade_options=GRADE_OPTIONS,
            letters=letters,
            rating_history=rating_history,
            semester_overview=semester_overview,
            trend_semesters=trend_semesters,
            competency_trends=competency_trends,
            semester_goal_text=semester_goal_text,
        )

    @app.route("/students/<int:student_id>/letters/generate", methods=["POST"])
    def generate_letter(student_id: int) -> Any:
        semester = request.form.get("semester", DEFAULT_SEMESTER)
        if semester in get_archived_semesters():
            flash("Für archivierte Halbjahre können keine neuen Lernbriefe erzeugt werden.", "error")
            return redirect(url_for("ratings", student_id=student_id, semester=semester))

        active_tpl = get_active_letter_template()

        try:
            content = build_letter(student_id, semester)
        except ValueError as exc:
            flash(str(exc), "error")
            return redirect(url_for("ratings", student_id=student_id, semester=semester))

        get_db().execute(
            """
            INSERT INTO letters (student_id, semester, content, created_at, template_name, body_font_family, body_font_size)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                student_id,
                semester,
                content,
                datetime.now().isoformat(timespec="seconds"),
                active_tpl["name"],
                active_tpl["body_font_family"],
                int(active_tpl["body_font_size"]),
            ),
        )
        get_db().commit()

        flash("Lernbrief wurde generiert und gespeichert.", "success")
        return redirect(url_for("ratings", student_id=student_id, semester=semester))

    @app.route("/letters/<int:letter_id>")
    def letter_detail(letter_id: int) -> Any:
        letter = get_db().execute(
            """
            SELECT l.id, l.student_id, l.content, l.semester, l.created_at,
                   l.template_name, l.body_font_family, l.body_font_size,
                   s.full_name
            FROM letters l
            JOIN students s ON s.id = l.student_id
            WHERE l.id = ?
            """,
            (letter_id,),
        ).fetchone()

        if letter is None:
            flash("Lernbrief nicht gefunden.", "error")
            return redirect(url_for("index"))

        return render_template("letter_detail.html", letter=letter)

    @app.route("/letters/<int:letter_id>/update", methods=["POST"])
    def update_letter(letter_id: int) -> Any:
        db = get_db()
        letter = db.execute(
            """
            SELECT l.id, l.student_id, l.semester
            FROM letters l
            WHERE l.id = ?
            """,
            (letter_id,),
        ).fetchone()

        if letter is None:
            flash("Lernbrief nicht gefunden.", "error")
            return redirect(url_for("index"))

        content = request.form.get("content_html", "").strip()
        if not content:
            flash("Lernbriefinhalt darf nicht leer sein.", "error")
            return redirect(url_for("letter_detail", letter_id=letter_id))

        db.execute("UPDATE letters SET content = ? WHERE id = ?", (content, letter_id))
        db.commit()
        flash("Lernbrief wurde gespeichert.", "success")
        return redirect(url_for("letter_detail", letter_id=letter_id))

    @app.route("/letters/<int:letter_id>/export/pdf")
    def export_letter_pdf(letter_id: int) -> Any:
        letter = get_db().execute(
            """
            SELECT l.id, l.student_id, l.content, l.semester, l.created_at,
                   l.body_font_family, l.body_font_size,
                   s.full_name
            FROM letters l
            JOIN students s ON s.id = l.student_id
            WHERE l.id = ?
            """,
            (letter_id,),
        ).fetchone()

        if letter is None:
            flash("Lernbrief nicht gefunden.", "error")
            return redirect(url_for("index"))

        pdf_buffer = render_letter_pdf(letter)
        filename = make_export_filename(letter["full_name"], letter["semester"], "pdf")
        return send_file(pdf_buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")

    @app.route("/letters/<int:letter_id>/export/docx")
    def export_letter_docx(letter_id: int) -> Any:
        letter = get_db().execute(
            """
            SELECT l.id, l.student_id, l.content, l.semester, l.created_at,
                   l.body_font_family, l.body_font_size,
                   s.full_name
            FROM letters l
            JOIN students s ON s.id = l.student_id
            WHERE l.id = ?
            """,
            (letter_id,),
        ).fetchone()

        if letter is None:
            flash("Lernbrief nicht gefunden.", "error")
            return redirect(url_for("index"))

        docx_buffer = render_letter_docx(letter)
        filename = make_export_filename(letter["full_name"], letter["semester"], "docx")
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    @app.route("/letters/<int:letter_id>/delete", methods=["POST"])
    def delete_letter(letter_id: int) -> Any:
        db = get_db()
        letter = db.execute(
            "SELECT id, student_id, semester FROM letters WHERE id = ?",
            (letter_id,),
        ).fetchone()

        if letter is None:
            flash("Lernbrief nicht gefunden.", "error")
            return redirect(url_for("index"))

        db.execute("DELETE FROM letters WHERE id = ?", (letter_id,))
        db.commit()
        flash("Lernbrief wurde gelöscht.", "success")

        next_target = request.form.get("next", "index")
        if next_target == "ratings":
            return redirect(
                url_for(
                    "ratings",
                    student_id=letter["student_id"],
                    semester=letter["semester"],
                )
            )

        return redirect(url_for("index"))

    return app


if __name__ == "__main__":
    app = create_app()
    host = os.getenv("LERNBRIEF_HOST", "127.0.0.1")
    port = int(os.getenv("LERNBRIEF_PORT", "5000"))
    debug_mode = os.getenv("FLASK_DEBUG", "0") == "1"

    should_open_browser = os.getenv("LERNBRIEF_OPEN_BROWSER", "1") == "1"
    is_reloader_child = os.getenv("WERKZEUG_RUN_MAIN") == "true"
    if should_open_browser and (not debug_mode or is_reloader_child):
        url = f"http://{host}:{port}"
        threading.Thread(
            target=open_browser_when_ready,
            args=(url, host, port),
            daemon=True,
        ).start()

    app.run(host=host, port=port, debug=debug_mode, use_reloader=debug_mode)
